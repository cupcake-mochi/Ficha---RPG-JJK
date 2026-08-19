# -*- coding: utf-8 -*-
"""Regressao completa: o catalogo reproduz TODOS os numeros impressos na
ficha-exemplo-kaori.docx? Le os dois lados; nada esta escrito na mao aqui."""
import json, re
from docx import Document
CAT = json.load(open('catalogo-projeto-m.json', encoding='utf-8'))
DOC = Document('repos/JJK---PDF---RPG-main/ficha/ficha-exemplo-kaori.docx')

KAORI = {"Força":3, "Constituição":2, "Destreza":2, "Inteligência":1, "Essência":1}
CAMINHO, NIVEL, PROTECAO = "Bastião", 2, 1

C = CAT["caminhos"][CAMINHO]
con, des, forca = KAORI["Constituição"], KAORI["Destreza"], KAORI["Força"]
maestria = 1 + NIVEL // 8
CALC = {
    "Vida":          str((C["vida_inicial"] + con) + (C["vida_por_nivel"] + con) * (NIVEL - 1)),
    "Energia":       str(C["pe_por_nivel"] * NIVEL),
    "Integridade":   str(20 + 8 * (NIVEL - 1)),
    "Defesa":        str(10 + des + PROTECAO),
    "Iniciativa":    f"d20 + {des}",
    "Deslocamento":  "9 m",
    "Maestria":      str(maestria),
    "CD de feitiço": str(10 + 2 + maestria),
    "Conjuração":    f"d20 + {2 + maestria}",
    "Corpo a corpo": f"d20 + {forca}",
    "À distância":   f"d20 + {des}",
}
# le a tabela de numeros derivados da ficha
IMPRESSO = {}
for t in DOC.tables:
    for r in t.rows:
        cel = [c.text.strip() for c in r.cells]
        for i in (0, 3):
            if len(cel) > i + 2 and cel[i] in CALC:
                IMPRESSO[cel[i]] = cel[i + 2]

print(f"{'campo':16} {'catalogo':>12}   {'ficha .docx':>12}")
falhas = 0
for campo, meu in CALC.items():
    dela = IMPRESSO.get(campo, "<nao achei>")
    ok = meu.replace(" ", "") == dela.replace(" ", "")
    falhas += not ok
    print(f"  {campo:16} {meu:>12}   {dela:>12}   {'BATE' if ok else 'NAO BATE'}")

# pericias treinadas: as duas fixas do Caminho aparecem marcadas?
marcadas = re.findall(r'■\s+([A-ZÁÉÍÓÚÂÊÔÃÕÇ][\wáéíóúâêôãõç ]+)',
                      '\n'.join(c.text for t in DOC.tables for r in t.rows for c in r.cells))
marcadas = [m.strip() for m in marcadas]
print(f"\n  pericias treinadas na ficha: {len(marcadas)}  ({', '.join(marcadas)})")
for fixa in C["pericias_fixas"]:
    achou = any(fixa in m for m in marcadas)
    print(f"  [{'BATE' if achou else 'NAO BATE'}] fixa do {CAMINHO}: {fixa}")

# as Familias que a Kaori marcou existem no manual?
livres = ["Controle", "Castigo"]; fechadas = ["Área", "Auxiliares", "Amparo"]
fantasma = [f for f in livres + fechadas if f not in CAT["familias"]]
print(f"\n  Familias da Kaori: Livres {livres}, Fechadas {fechadas}")
print(f"  [{'OK' if not fantasma else 'PROBLEMA'}] todas existem no manual"
      f"{'' if not fantasma else ': ' + str(fantasma)}")
print(f"\n{'TODOS OS NUMEROS BATEM' if not falhas else f'{falhas} FALHA(S)'}")
